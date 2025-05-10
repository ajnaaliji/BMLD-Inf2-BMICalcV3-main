import streamlit as st
import pandas as pd
import io
import base64
from datetime import datetime
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
import tempfile
from PIL import Image
from docx.shared import Inches
import uuid
from utils.data_manager import DataManager
from utils.ui_helpers import apply_theme

# ==== Initialisierung ====
st.set_page_config(page_title="Klinische Chemie", page_icon="🧪")
apply_theme()
data_manager = DataManager()
username = st.session_state.get("username", "anonymous")
data_manager.load_user_data("klinische_eintraege", f"data_klinische_chemie_{username}.csv", initial_value=[])

# ==== DataHandler vorbereiten ====
dh_word = data_manager._get_data_handler(f"word_klinische_chemie/{username}")
dh_img = data_manager._get_data_handler(f"bilder_klinische_chemie/{username}")
dh_docs = data_manager._get_data_handler(f"anhang_klinische_chemie/{username}")
for dh in [dh_word, dh_img, dh_docs]:
    if not dh.filesystem.exists(dh.root_path):
        dh.filesystem.makedirs(dh.root_path)

# ==== Icon laden ====
def load_icon_base64(path):
    with open(path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode("utf-8")

img_icon = load_icon_base64(
    "assets/clinical_chemistry.png")
img_paper = load_icon_base64("assets/paperclip.png")
img_pic = load_icon_base64("assets/picture.png")
img_post = load_icon_base64("assets/postanalyt.png")
img_sample = load_icon_base64("assets/medical-sample.png")
img_monitor = load_icon_base64("assets/monitor.png")
img_befund = load_icon_base64("assets/befunde.png")
img_anamnese = load_icon_base64("assets/anamnese.png")
img_angabe = load_icon_base64("assets/angabe.png")
img_patient = load_icon_base64("assets/patient.png")
img_stamp = load_icon_base64("assets/stamp.png")
img_kontrolle = load_icon_base64("assets/medicine.png")

# ==== Titel ====
st.markdown(f"""
<h1 style='display: flex; align-items: center; gap: 24px;'>
    Klinische Chemie
    <img src='data:image/png;base64,{img_icon}' width='50'>
</h1>
""", unsafe_allow_html=True)

# ==== Patientenangaben ====
st.markdown(f"""
<div style="display: flex; align-items: center; margin-bottom: 1rem;">
    <img src="data:image/png;base64,{img_angabe}" style="height: 40px; margin-right: 10px;" />
    <h2 style="margin: 0;">Patientenangaben</h2>
</div>
""", unsafe_allow_html=True)

col1, col2 = st.columns(2)
with col1:
    patient_name = st.text_input("Patientenname")
    geburtstag = st.text_input("Geburtstag/Alter")
    geschlecht = st.selectbox("Geschlecht", ["", "weiblich", "männlich", "divers"])
    semester = st.selectbox("Semester", ["1", "2", "3", "4", "5", "6"], key="semester")
with col2:
    groesse = st.text_input("Grösse (cm)")
    gewicht = st.text_input("Gewicht (kg)")
    datum = st.date_input("Befunddatum", value=datetime.today())

# ==== Anamnese ====
st.markdown(f"""
<div style="display: flex; align-items: center; margin-bottom: 1rem;">
    <img src="data:image/png;base64,{img_anamnese}" style="height: 40px; margin-right: 10px;" />
    <h2 style="margin: 0;">Anamnese</h2>
</div>
""", unsafe_allow_html=True)
anamnese = st.text_area("Anamnese", height=100)

# ==== Vorbefunde ====
st.markdown(f"""
<div style="display: flex; align-items: center; margin-bottom: 1rem;">
    <img src="data:image/png;base64,{img_befund}" style="height: 40px; margin-right: 10px;" />
    <h2 style="margin: 0;">Vorbefunde</h2>
</div>
""", unsafe_allow_html=True)
vorbefunde = st.text_area("Vorbefunde", height=100)

# ==== Präanalytik ====
st.markdown(f"""
<div style="display: flex; align-items: center; margin-bottom: 1rem;">
    <img src="data:image/png;base64,{img_sample}" style="height: 40px; margin-right: 10px;" />
    <h2 style="margin: 0;">Präanalytik</h2>
</div>
""", unsafe_allow_html=True)
probenmaterial = st.text_input("Probenmaterial")
makro = st.text_area("Makroskopische Beurteilung", height=100)

# ==== Analytik ====
st.markdown(f"""
<div style="display: flex; align-items: center; margin-bottom: 1rem;">
    <img src="data:image/png;base64,{img_monitor}" style="height: 40px; margin-right: 10px;" />
    <h2 style="margin: 0;">Analytik</h2>
</div>
""", unsafe_allow_html=True)
reagenzien = st.text_area("Reagenzien (Name/LOT/Verfall)", height=100)
qc = st.text_area("Qualitätskontrolle (Name/LOT/Verfall)", height=100)
methode = st.text_input("Analyt/Methode/Gerät")
validation = st.text_area("Technische Validation der QC (Sollwertbereich/Ist-gemessener QK-Wert)", height=100)

# ==== Postanalytik ====
st.markdown(f"""
<div style="display: flex; align-items: center; margin-bottom: 1rem;">
    <img src="data:image/png;base64,{img_post}" style="height: 40px; margin-right: 10px;" />
    <h2 style="margin: 0;">Postanalytik</h2>
</div>
""", unsafe_allow_html=True)
st.markdown("<h4 style='margin-top: 1rem; margin-bottom: 0.3rem; text-decoration: underline; color: #333;'>Biomedizinische Validation</h4>", unsafe_allow_html=True)
transversal = st.text_area("Transversalbeurteilung", height=100)
plausi = st.text_area("Plausibilitätskontrolle", height=100)
st.markdown(f"""
<div style="display: flex; align-items: center; margin-bottom: 0.8rem;">
    <img src="data:image/png;base64,{img_kontrolle}" style="height: 40px; margin-right: 10px;" />
    <h3 style="margin: 0;">Plausibilitätskontrolle</h3>
</div>
""", unsafe_allow_html=True)
extremwerte = st.text_area("Extremwerte", height=100)
konstellation = st.text_area("Konstellationskontrolle der Werte miteinander", height=100)

# ==== Freigabe ====
st.markdown(f"""
<div style="display: flex; align-items: center; margin-bottom: 1rem;">
    <img src="data:image/png;base64,{img_stamp}" style="height: 40px; margin-right: 10px;" />
    <h2 style="margin: 0;">Freigabeentscheidung</h2>
</div>
""", unsafe_allow_html=True)
freigabe = st.text_area("Freigabeentscheidung", height=100)

trend = st.text_area("Trend (zu Vorbefunden, falls vorhanden)", height=100)
konstellation = st.text_area("Konstellationskontrolle der Werte miteinander", height=100)

felder = {
    "vorbefunde": vorbefunde,
    "probenmaterial": probenmaterial,
    "makro": makro,
    "reagenzien": reagenzien,
    "qc": qc,
    "methode": methode,
    "validation": validation,
    "transversal": transversal,
    "plausi": plausi,
    "extremwerte": extremwerte,
    "trend": trend,
    "konstellation": konstellation,
}

# ==== Bilder uploaden ====
st.markdown(f"""
<h3 style='display: flex; align-items: center; gap: 10px;'>
    <img src='data:image/png;base64,{img_pic}' width='40'>
    Mikroskopiebilder oder Befundbilder
</h3>
""", unsafe_allow_html=True)
uploaded_images = st.file_uploader("Bilder auswählen", type=["png", "jpg", "jpeg"], accept_multiple_files=True)
temp_uploaded_images = []

if uploaded_images:
    st.markdown("**Vorschau:**")
    bestehende_bilder = [f["name"] for f in dh_img.filesystem.ls(dh_img.root_path)]
    for img in uploaded_images:
        st.image(img, use_container_width=True)
        name = img.name.replace(" ", "_")
        if name in bestehende_bilder:
            st.info(f"⏭️ Bild bereits vorhanden: {name}")
            continue
        try:
            image_pil = Image.open(img)
            image_pil.verify()
            img.seek(0)
            image_converted = Image.open(img).convert("RGB")
            temp_bytes = io.BytesIO()
            image_converted.save(temp_bytes, format="JPEG")
            temp_bytes.seek(0)
            temp_uploaded_images.append((name, temp_bytes.getvalue()))
        except Exception as e:
            st.warning(f"⚠️ Bild beschädigt oder ungeeignet: {name} ({e})")

# ==== Anhänge uploaden ====
st.markdown(f"""
<h3 style='display: flex; align-items: center; gap: 10px;'>
    <img src='data:image/png;base64,{img_paper}' width='40'>
    Weitere Dateien (PDF, Word)
</h3>
""", unsafe_allow_html=True)
uploaded_docs = st.file_uploader("Dateien auswählen", type=["pdf", "docx"], accept_multiple_files=True)
temp_uploads = [(f.name.replace(" ", "_"), f.getvalue()) for f in uploaded_docs] if uploaded_docs else []
anhang_dateien = []

# ==== Speichern und Exportieren ====
if st.button("📁 Speichern und Exportieren"):
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")

    # Bilder speichern
    for name, img_bytes in temp_uploaded_images:
        dh_img.save(f"{timestamp}_{uuid.uuid4().hex}_{name}", img_bytes)

    # Anhänge speichern
    for name, content in temp_uploads:
        fname = f"{timestamp}_{uuid.uuid4().hex[:8]}_{name}"
        dh_docs.save(fname, content)
        anhang_dateien.append(fname)

    # Word-Datei erstellen
    doc = Document()
    doc.add_heading("Klinische Chemie", 0)
    doc.add_paragraph(f"Bericht vom {datum.strftime('%d.%m.%Y')} – Patient: {patient_name}")
    doc.add_paragraph("")  # Leerzeile

    # Patientenangaben
    doc.add_heading("Patientenangaben", level=1)
    for label, value in [
        ("Patientenname", patient_name),
        ("Geburtstag/Alter", geburtstag),
        ("Geschlecht", geschlecht),
        ("Größe (cm)", groesse),
        ("Gewicht (kg)", gewicht)
    ]:
        doc.add_heading(label, level=2)
        doc.add_paragraph(value)

    # Weitere Abschnitte
    abschnitte = {
        "Anamnese": {"Anamnese": anamnese},
        "Präanalytik": {
            "Probenmaterial": probenmaterial,
            "Makroskopische Beurteilung": makro
        },
        "Analytik": {
            "Reagenzien": reagenzien,
            "Qualitätskontrolle": qc,
            "Analyt/Methode/Gerät": methode,
            "Technische Validation": validation
        },
        "Postanalytik": {
            "Transversalbeurteilung": transversal,
                "Plausibilitätskontrolle": plausi,
            "Extremwerte": extremwerte,
            "Konstellation": konstellation
        },
        "Freigabeentscheidung": {
            "Freigabe": freigabe
        },
        "Weitere Angaben": {
            "Vorbefunde": felder.get("vorbefunde", ""),
            "Trend": felder.get("trend", "")
        }
    }

    # Abschnitte in Word einfügen
    for abschnitt, inhalte in abschnitte.items():
        doc.add_heading(abschnitt, level=1)
        for label, value in inhalte.items():
            doc.add_heading(label, level=2)
            doc.add_paragraph(str(value))

    if temp_uploaded_images:
        doc.add_page_break()
        doc.add_heading("Mikroskopiebilder / Befundbilder", level=1)
        for name, img_bytes in temp_uploaded_images:
            try:
                img_stream = io.BytesIO(img_bytes)
                doc.add_picture(img_stream, width=Inches(5))  # Breite nach Wunsch
                doc.add_paragraph(name)
            except Exception as e:
                doc.add_paragraph(f"⚠️ Bild konnte nicht eingefügt werden: {name} ({e})")

    # Word speichern
    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    filename_word = f"{timestamp}_{uuid.uuid4().hex[:6]}_bericht.docx"
    dh_word.save(filename_word, word_buffer.getvalue())

    # PDF erstellen
    def add_section(c, title, fields, x, y):
        c.setFont("Helvetica-Bold", 14)
        c.drawString(x, y, title)
        y -= 0.8 * cm
        c.setFont("Helvetica", 12)
        for label, value in fields.items():
           for line in str(value).splitlines():
                c.drawString(x, y, f"{label}: {line}" if label else line)
                y -= 0.6 * cm
                label = ""
                if y < 2 * cm:
                    c.showPage()
                    y = A4[1] - 2 * cm
        y -= 0.5 * cm
        return y

    # PDF generieren
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
        c = canvas.Canvas(tmp_pdf.name, pagesize=A4)
        x, y = 2 * cm, A4[1] - 2 * cm

        # Kopf
        c.setFont("Helvetica-Bold", 16)
        c.drawString(x, y, "Klinische Chemie")
        y -= 1 * cm
        c.setFont("Helvetica", 12)
        c.drawString(x, y, f"Bericht vom {datum.strftime('%d.%m.%Y')} – Patient: {patient_name}")
        y -= 1.5 * cm

        # Abschnitte
        y = add_section(c, "Patientenangaben", {
            "Patientenname": patient_name,
            "Geburtstag/Alter": geburtstag,
            "Geschlecht": geschlecht,
            "Größe (cm)": groesse,
            "Gewicht (kg)": gewicht
        }, x, y)

        y = add_section(c, "Anamnese", {"Anamnese": anamnese}, x, y)
        y = add_section(c, "Vorbefunde", {"Vorbefunde": felder.get("vorbefunde", "")}, x, y)
        y = add_section(c, "Präanalytik", {
            "Probenmaterial": probenmaterial,
            "Makroskopische Beurteilung": makro
        }, x, y)
        y = add_section(c, "Analytik", {
            "Reagenzien": reagenzien,
            "Qualitätskontrolle": qc,
            "Analyt/Methode/Gerät": methode,
            "Technische Validation": validation
        }, x, y)
        y = add_section(c, "Postanalytik", {
            "Transversalbeurteilung": transversal,
            "Plausibilitätskontrolle": plausi,
            "Extremwerte": extremwerte,
            "Konstellation": konstellation
        }, x, y)
        y = add_section(c, "Trend", {"Trend": felder.get("trend", "")}, x, y)
        y = add_section(c, "Freigabeentscheidung", {"Freigabe": freigabe}, x, y)

        # Bilder (optional)
        if temp_uploaded_images:
            c.showPage()
            y = A4[1] - 2 * cm
            c.setFont("Helvetica-Bold", 14)
            c.drawString(x, y, "🖼 Mikroskopiebilder / Befundbilder")
            y -= 1.2 * cm

            for name, img_bytes in temp_uploaded_images:
                try:
                    image = Image.open(io.BytesIO(img_bytes)).convert("RGB")
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_img:
                        image.save(tmp_img.name, format="PNG")
                        c.drawImage(tmp_img.name, x, y - 6 * cm, width=12 * cm, height=6 * cm)
                        y -= 7 * cm
                        if y < 3 * cm:
                            c.showPage()
                            y = A4[1] - 2 * cm
                except Exception as e:
                    st.warning(f"⚠️ Bild konnte nicht eingefügt werden: {name} ({e})")

        # Speichern (immer nur hier!)
        c.save()
        pdf_filename = f"{timestamp}_{uuid.uuid4().hex[:6]}_bericht.pdf"
        with open(tmp_pdf.name, "rb") as f:
            pdf_bytes = f.read()
            dh_docs.save(pdf_filename, pdf_bytes)
            anhang_dateien.append(pdf_filename)

    # Speichern in CSV
    neuer_eintrag = {
        "titel": patient_name if patient_name.strip() else f"Bericht vom {datum.strftime('%d.%m.%Y')}",
        "zeit": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "patient_name": patient_name,
        "geburtstag": geburtstag,
        "geschlecht": geschlecht,
        "groesse": groesse,
        "gewicht": gewicht,
        "datum": datum.strftime("%Y-%m-%d"),
        "anhaenge": anhang_dateien,
        "semester": str(semester),
        "dateiname": filename_word,
        **felder
    }

    df_neu = pd.DataFrame([neuer_eintrag])
    if isinstance(st.session_state["klinische_eintraege"], pd.DataFrame):
        st.session_state["klinische_eintraege"] = pd.concat([
            st.session_state["klinische_eintraege"], df_neu
        ], ignore_index=True)
    else:
        st.session_state["klinische_eintraege"] = df_neu

    data_manager.save_data("klinische_eintraege")
    st.success("✅ Eintrag gespeichert!")

    st.download_button("⬇️ Word herunterladen", word_buffer, file_name=filename_word, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    st.download_button("⬇️ PDF herunterladen", pdf_bytes, file_name=pdf_filename, mime="application/pdf")

    st.markdown("📎 Zum Teilen einfach herunterladen und weitergeben.")

# ==== Zurück ==== 
if st.button("🔙 Zurück zur Übersicht"):
    st.switch_page("pages/01_Datei.py")
