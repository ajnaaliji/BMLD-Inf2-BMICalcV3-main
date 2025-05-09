import streamlit as st
import os
import yaml
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches
from PIL import Image
import tempfile
from utils.data_manager import DataManager
from zoneinfo import ZoneInfo
from os.path import basename
from utils.ui_helpers import apply_theme
import base64

# ==== Icon laden ====
def load_icon_base64(path):
    with open(path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode("utf-8")
    
img_safe = load_icon_base64("assets/security.png")
img_guide = load_icon_base64("assets/guideline.png")
img_pic = load_icon_base64("assets/picture.png")
img_load = load_icon_base64("assets/download.png")

# === Setup ===
st.set_page_config(
    page_title="Zellatlas H√§matologie",
    page_icon="assets/guideline.png",  # ‚Üê Das ist der relative Pfad zur Datei
    layout="centered",
    initial_sidebar_state="collapsed"
)
apply_theme()
if "username" not in st.session_state or st.session_state["authentication_status"] is not True:
    st.error("‚ùå Nicht eingeloggt ‚Äì bitte zuerst anmelden.")
    st.stop()

username = st.session_state["username"]
data_manager = DataManager()
atlas_folder = f"zellatlas_haematologie/{username}"
dh = data_manager._get_data_handler(atlas_folder)
if not dh.filesystem.exists(dh.root_path):
    dh.filesystem.makedirs(dh.root_path)

# === Kategorien ===
bereiche = {
    "Wei√ües Blutbild": [
        "Myeloblast", "Promyelozyt", "Myelozyt", "Metamyelozyt",
        "Stabkerniger", "Segmentkerniger", "Eosinophiler", "Basophiler",
        "Monozyt", "Lymphozyt", "Plasmazelle", "Nicht klassifizierbar (wei√ü)"
    ],
    "Rotes Blutbild": [
        "Normozyt", "Mikrozyt", "Makrozyt", "Fragmentozyt",
        "Targetzelle", "Sichelzelle", "Nicht klassifizierbar (rot)"
    ],
    "Thrombozyten": [
        "Normal", "Riesenthrombozyt", "Agranul√§r", "Nicht klassifizierbar (thrombo)"
    ]
}

# === Eintragsliste initialisieren ===
if "zell_eintraege" not in st.session_state:
    st.session_state.zell_eintraege = [{}]

# === Eingabeformulare ===
st.markdown(f"""
<h1 style='display: flex; align-items: center; gap: 16px; margin-bottom: 0;'>
    <img src='data:image/png;base64,{img_guide}' width='48'>
    Zellatlas H√§matologie
</h1>
<p style='margin-top: 4px; font-size: 20px; text-decoration: underline; color: #333;'>
    Zell-Eintr√§ge erfassen
</p>
""", unsafe_allow_html=True)

for idx, eintrag in enumerate(st.session_state.zell_eintraege):
    st.markdown(f"## Eintrag {idx + 1}")
    typ = st.selectbox("Zelltyp w√§hlen", [f"{k}: {v}" for k in bereiche for v in bereiche[k]], key=f"typ_{idx}")
    st.markdown(f"""
    <h4 style='display: flex; align-items: center; gap: 10px; margin-top: 30px;'>
        <img src='data:image/png;base64,{img_pic}' width='34'>
        Bild hochladen (png/jpg)
    </h4>
    """, unsafe_allow_html=True)

bild = st.file_uploader("", type=["png", "jpg", "jpeg"], key=f"bild_{idx}")
beschreibung = st.text_area("Beschreibung / Merkmale", key=f"beschreibung_{idx}")
st.session_state.zell_eintraege[idx] = {"typ": typ, "beschreibung": beschreibung, "bild": bild}

# === Plus-Button ===
st.markdown("---")
if st.button("‚ûï Weiteren Eintrag hinzuf√ºgen"):
    st.session_state.zell_eintraege.append({})
    st.rerun()

# === Speichern ===
if st.button("üìÇ Alle Eintr√§ge speichern"):
    erfolgreich = False
    for eintrag in st.session_state.zell_eintraege:
        if eintrag.get("typ") and eintrag.get("beschreibung"):
            now = datetime.now(ZoneInfo("Europe/Zurich"))
            timestamp = now.strftime("%Y%m%dT%H%M%S")
            eintrag_data = {
                "typ": eintrag["typ"],
                "beschreibung": eintrag["beschreibung"],
                "zeit": now.isoformat()
            }
            if eintrag["bild"]:
                img_bytes = eintrag["bild"].getvalue()
                img_name = f"{timestamp}_{eintrag['bild'].name.replace(' ', '_')}"
                dh.save(img_name, img_bytes)
                eintrag_data["bild"] = img_name

            yaml_name = f"{timestamp}_{eintrag['typ'].split(':')[1].strip().lower().replace(' ', '_')}.yaml"
            dh.write_text(yaml_name, yaml.dump(eintrag_data, allow_unicode=True))
            erfolgreich = True

    if erfolgreich:
        st.success("‚úÖ Eintr√§ge wurden gespeichert.")
        st.session_state.zell_eintraege = [{}]
    else:
        st.warning("‚ö†Ô∏è Keine g√ºltigen Eintr√§ge zum Speichern gefunden.")

# === Anzeige gespeicherter Eintr√§ge ===
st.markdown(f"""
<h2 style='display: flex; align-items: center; gap: 10px;'>
    <img src='data:image/png;base64,{img_safe}' width='50'>
    Gespeicherte Zell-Eintr√§ge
</h2>
""", unsafe_allow_html=True)
eintrags_liste = [
    f["name"] if isinstance(f, dict) else f
    for f in dh.filesystem.ls(dh.root_path)
    if (f["name"] if isinstance(f, dict) else f).endswith(".yaml")
]

if eintrags_liste:
    eintrags_liste.sort(reverse=True)
    for filename in eintrags_liste:
        try:
            data = yaml.safe_load(dh.read_text(basename(filename)))
            st.markdown(f"### üî¨ {data['typ']}")
            zeit_raw = data.get("zeit", "")
            try:
                zeit_formatiert = datetime.fromisoformat(zeit_raw).strftime("%d.%m.%Y, %H:%M Uhr")
                st.markdown(f"{zeit_formatiert}")
            except:
                st.markdown(f"{zeit_raw}")
            st.markdown(data.get("beschreibung", "Keine Beschreibung vorhanden."))
            if "bild" in data:
                st.image(dh.read_binary(basename(data["bild"])), width=300)

            # === L√∂sch-Logik mit einfacher Best√§tigung ===
            if st.button(f"üóëÔ∏è L√∂schen", key=f"delete_{filename}"):
                st.session_state[f"confirm_delete_{filename}"] = True

            if st.session_state.get(f"confirm_delete_{filename}", False):
                st.warning(f"M√∂chtest du **{data['typ']}** wirklich l√∂schen?")
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("‚ùå Nein", key=f"cancel_{filename}"):
                        st.session_state[f"confirm_delete_{filename}"] = False
                        st.rerun()
                with col2:
                    if st.button("‚úÖ Ja", key=f"confirm_{filename}"):
                        try:
                            dh.filesystem.delete(os.path.join(dh.root_path, basename(filename)))
                            if "bild" in data:
                                bild_pfad = os.path.join(dh.root_path, basename(data["bild"]))
                                if dh.filesystem.exists(bild_pfad):
                                    dh.filesystem.delete(bild_pfad)
                            st.success("‚úÖ Eintrag gel√∂scht.")
                            st.rerun()
                        except Exception as e:
                            st.warning(f"‚ö†Ô∏è Fehler beim L√∂schen von {filename}: {e}")

            st.markdown("---")
        except Exception as e:
            st.warning(f"‚ö†Ô∏è Fehler beim Laden von {filename}: {e}")
else:
    st.info("Noch keine Eintr√§ge vorhanden.")

# === Neuesten Eintrag exportieren ===
if eintrags_liste:
    neueste_datei = sorted(eintrags_liste, reverse=True)[0]
    data = yaml.safe_load(dh.read_text(basename(neueste_datei)))
    typ = data.get("typ", "Unbekannt")
    beschreibung = data.get("beschreibung", "")
    zeit_raw = data.get("zeit", "")
    zeit_dt = datetime.fromisoformat(zeit_raw) if zeit_raw else datetime.now()
    zeit_fmt = zeit_dt.strftime("%Y-%m-%d %H:%M")

    # Word exportieren
    doc = Document()
    doc.add_heading(f"Zellatlas-Eintrag: {typ}", 0)
    doc.add_paragraph(f"Erstellt am: {zeit_fmt}")
    doc.add_paragraph(beschreibung)
    if "bild" in data:
        try:
            image_data = dh.read_binary(data["bild"])
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_img:
                tmp_img.write(image_data)
                doc.add_picture(tmp_img.name, width=Inches(4.5))
        except:
            st.warning("‚ö†Ô∏è Bild konnte nicht in Word eingef√ºgt werden.")

    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)

    # PDF exportieren
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm

    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
        c = canvas.Canvas(tmp_pdf.name, pagesize=A4)
        x, y = 2 * cm, A4[1] - 2 * cm
        c.setFont("Helvetica-Bold", 16)
        c.drawString(x, y, f"Zellatlas-Eintrag: {typ}")
        y -= 1.5 * cm
        c.setFont("Helvetica", 12)
        c.drawString(x, y, f"Erstellt am: {zeit_fmt}")
        y -= 1.5 * cm
        for line in beschreibung.splitlines():
            c.drawString(x, y, line.strip())
            y -= 0.6 * cm
            if y < 2 * cm:
                c.showPage()
                y = A4[1] - 2 * cm
        c.save()
        with open(tmp_pdf.name, "rb") as f:
            pdf_data = f.read()

    # Download-Buttons anzeigen
    st.markdown("""
    <h2 style='text-decoration: underline; font-weight: bold; font-size: 28px; margin-top: 20px;'>
        Zell-Eintr√§ge erfassen
    </h2>
    """, unsafe_allow_html=True)

    st.download_button("‚¨áÔ∏è Word", data=word_buffer, file_name=f"Zellatlas_{typ}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    st.download_button("‚¨áÔ∏è PDF", data=pdf_data, file_name=f"Zellatlas_{typ}.pdf", mime="application/pdf")

# === Zur√ºck zur √úbersicht ===
st.markdown("---")
if st.button("üîô Zur√ºck zur √úbersicht"):
    st.switch_page("pages/01_Datei.py")
