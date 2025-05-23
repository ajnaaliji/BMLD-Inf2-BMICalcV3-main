import streamlit as st
import os
import yaml
import io
from datetime import datetime
from docx import Document
from docx.shared import Inches
from PIL import Image
import tempfile
from utils.data_manager import DataManager
from zoneinfo import ZoneInfo
from os.path import basename
from utils.ui_helpers import apply_theme
import base64

# ==== Icon laden ====
def load_icon_base64(path):
    with open(path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode("utf-8")
    
img_safe = load_icon_base64("assets/security.png")
img_guide = load_icon_base64("assets/guideline.png")
img_pic = load_icon_base64("assets/picture.png")
img_load = load_icon_base64("assets/download.png")

# === Setup ===
st.set_page_config(
    page_title="Zellatlas Hämatologie",
    page_icon="assets/guideline.png",  # ← Das ist der relative Pfad zur Datei
    layout="centered",
    initial_sidebar_state="collapsed"
)
apply_theme()
if "username" not in st.session_state or st.session_state["authentication_status"] is not True:
    st.error("❌ Nicht eingeloggt – bitte zuerst anmelden.")
    st.stop()

username = st.session_state["username"]
data_manager = DataManager()
atlas_folder = f"zellatlas_haematologie/{username}"
dh = data_manager._get_data_handler(atlas_folder)
if not dh.filesystem.exists(dh.root_path):
    dh.filesystem.makedirs(dh.root_path)

# === Kategorien ===
bereiche = {
    "Weisses Blutbild": [
        "Myeloblast", "Promyelozyt", "Myelozyt", "Metamyelozyt",
        "Stabkerniger", "Segmentkerniger", "Eosinophiler", "Basophiler",
        "Monozyt", "Lymphozyt", "Plasmazelle", "Nicht klassifizierbar (weiß)"
    ],
    "Rotes Blutbild": [
        "Normozyt", "Mikrozyt", "Makrozyt", "Fragmentozyt",
        "Targetzelle", "Sichelzelle", "Nicht klassifizierbar (rot)"
    ],
    "Thrombozyten": [
        "Normal", "Riesenthrombozyt", "Agranulär", "Nicht klassifizierbar (thrombo)"
    ]
}

# === Eintragsliste initialisieren ===
if "zell_eintraege" not in st.session_state:
    st.session_state.zell_eintraege = [{}]

# === Eingabeformulare ===
st.markdown(f"""
<h1 style='display: flex; align-items: center; gap: 16px; margin-bottom: 0;'>
    Zellatlas Hämatologie
    <img src='data:image/png;base64,{img_guide}' width='48'>
</h1>
<p style='margin-top: 4px; font-size: 20px; text-decoration: underline; color: #333;'>
    Zell-Einträge erfassen
</p>
""", unsafe_allow_html=True)

for idx, eintrag in enumerate(st.session_state.zell_eintraege):
    st.markdown(f"## Eintrag {idx + 1}")
    typ = st.selectbox("Zelltyp wählen", [f"{k}: {v}" for k in bereiche for v in bereiche[k]], key=f"typ_{idx}")
    st.markdown(f"""
    <h4 style='display: flex; align-items: center; gap: 10px; margin-top: 30px;'>
        <img src='data:image/png;base64,{img_pic}' width='34'>
        Bild hochladen (png/jpg)
    </h4>
    """, unsafe_allow_html=True)

upload = st.file_uploader("", type=["png", "jpg", "jpeg"], key=f"bild_{idx}")
bildname = upload.name if upload else None
if upload is not None:
    try:
        img = Image.open(upload)
        img.verify()
        upload.seek(0)
        img = Image.open(upload).convert("RGB")
        temp_bytes = io.BytesIO()
        img.save(temp_bytes, format="JPEG")
        temp_bytes.seek(0)
        bild = temp_bytes
    except Exception as e:
        st.warning(f"⚠️ Hochgeladenes Bild ist beschädigt oder ungeeignet: {e}")
        bild = None
else:
    bild = None

beschreibung = st.text_area("Beschreibung / Merkmale", key=f"beschreibung_{idx}")
st.session_state.zell_eintraege[idx] = {
    "typ": typ,
    "beschreibung": beschreibung,
    "bild": bild,
    "bildname": bildname
}

# === Plus-Button ===
st.markdown("---")
if st.button("➕ Weiteren Eintrag hinzufügen"):
    st.session_state.zell_eintraege.append({})
    st.rerun()

# === Speichern ===
if st.button("📂 Alle Einträge speichern"):
    erfolgreich = False
    for eintrag in st.session_state.zell_eintraege:
        if eintrag.get("typ") and eintrag.get("beschreibung"):
            now = datetime.now(ZoneInfo("Europe/Zurich"))
            timestamp = now.strftime("%Y%m%dT%H%M%S")
            eintrag_data = {
                "typ": eintrag["typ"],
                "beschreibung": eintrag["beschreibung"],
                "zeit": now.isoformat()
            }
            if eintrag["bild"]:
                img_bytes = eintrag["bild"].getvalue()
                bildname = eintrag.get("bildname", "bild.jpg").replace(" ", "_")
                img_name = f"{timestamp}_{bildname}"
                dh.save(img_name, img_bytes)
                eintrag_data["bild"] = img_name

            yaml_name = f"{timestamp}_{eintrag['typ'].split(':')[1].strip().lower().replace(' ', '_')}.yaml"
            dh.write_text(yaml_name, yaml.dump(eintrag_data, allow_unicode=True))
            erfolgreich = True

    if erfolgreich:
        st.success("✅ Einträge wurden gespeichert.")
        st.session_state.zell_eintraege = [{}]
    else:
        st.warning("⚠️ Keine gültigen Einträge zum Speichern gefunden.")

# === Anzeige gespeicherter Einträge ===
st.markdown(f"""
<h2 style='display: flex; align-items: center; gap: 10px;'>
    <img src='data:image/png;base64,{img_safe}' width='50'>
    Gespeicherte Zell-Einträge
</h2>
""", unsafe_allow_html=True)
eintrags_liste = [
    f["name"] if isinstance(f, dict) else f
    for f in dh.filesystem.ls(dh.root_path)
    if (f["name"] if isinstance(f, dict) else f).endswith(".yaml")
]

if eintrags_liste:
    eintrags_liste.sort(reverse=True)
    for filename in eintrags_liste:
        try:
            data = yaml.safe_load(dh.read_text(basename(filename)))
            st.markdown(f"### {data['typ']}")
            zeit_raw = data.get("zeit", "")
            try:
                zeit_formatiert = datetime.fromisoformat(zeit_raw).strftime("%d.%m.%Y, %H:%M Uhr")
                st.markdown(f"{zeit_formatiert}")
            except:
                st.markdown(f"{zeit_raw}")
            st.markdown(data.get("beschreibung", "Keine Beschreibung vorhanden."))
            if "bild" in data:
                st.image(dh.read_binary(basename(data["bild"])), width=300)

            # === Lösch-Logik mit einfacher Bestätigung ===
            if st.button(f"🗑️ Löschen", key=f"delete_{filename}"):
                st.session_state[f"confirm_delete_{filename}"] = True

            if st.session_state.get(f"confirm_delete_{filename}", False):
                st.warning(f"Möchtest du **{data['typ']}** wirklich löschen?")
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("❌ Nein", key=f"cancel_{filename}"):
                        st.session_state[f"confirm_delete_{filename}"] = False
                        st.rerun()
                with col2:
                    if st.button("✅ Ja", key=f"confirm_{filename}"):
                        try:
                            dh.filesystem.delete(os.path.join(dh.root_path, basename(filename)))
                            if "bild" in data:
                                bild_pfad = os.path.join(dh.root_path, basename(data["bild"]))
                                if dh.filesystem.exists(bild_pfad):
                                    dh.filesystem.delete(bild_pfad)
                            st.success("✅ Eintrag gelöscht.")
                            st.rerun()
                        except Exception as e:
                            st.warning(f"⚠️ Fehler beim Löschen von {filename}: {e}")

            st.markdown("---")
        except Exception as e:
            st.warning(f"⚠️ Fehler beim Laden von {filename}: {e}")
else:
    st.info("Noch keine Einträge vorhanden.")

# === Alle Einträge gemeinsam exportieren ===
if eintrags_liste:
    eintrags_liste.sort()  # Chronologische Reihenfolge

    # === Word-Datei erstellen ===
    doc = Document()
    doc.add_heading("Zellatlas Hämatologie", 0)

    for datei in eintrags_liste:
        data = yaml.safe_load(dh.read_text(basename(datei)))
        typ = data.get("typ", "Unbekannt")
        beschreibung = data.get("beschreibung", "")
        zeit_raw = data.get("zeit", "")
        zeit_dt = datetime.fromisoformat(zeit_raw) if zeit_raw else datetime.now()
        zeit_fmt = zeit_dt.strftime("%d.%m.%Y %H:%M")

        doc.add_heading(f"{typ}", level=1)
        doc.add_paragraph(f"Erstellt am: {zeit_fmt}")
        doc.add_paragraph(beschreibung)

        if "bild" in data:
            try:
                image_data = dh.read_binary(data["bild"])
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_img:
                    tmp_img.write(image_data)
                    doc.add_picture(tmp_img.name, width=Inches(4.5))
            except Exception as e:
                doc.add_paragraph(f"⚠️ Bild {data.get('bild', 'unbekannt')} konnte nicht eingefügt werden: {e}")

        doc.add_paragraph("---")

    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm

with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
    c = canvas.Canvas(tmp_pdf.name, pagesize=A4)
    x, y = 2 * cm, A4[1] - 2 * cm

    for datei in eintrags_liste[::-1]:  # ältester unten, neuester oben
        data = yaml.safe_load(dh.read_text(basename(datei)))
        typ = data.get("typ", "Unbekannt")
        beschreibung = data.get("beschreibung", "")
        zeit_raw = data.get("zeit", "")
        zeit_dt = datetime.fromisoformat(zeit_raw) if zeit_raw else datetime.now()
        zeit_fmt = zeit_dt.strftime("%d.%m.%Y %H:%M")

        c.setFont("Helvetica-Bold", 14)
        c.drawString(x, y, f"Zelltyp: {typ}")
        y -= 1 * cm
        c.setFont("Helvetica", 12)
        c.drawString(x, y, f"Erstellt am: {zeit_fmt}")
        y -= 1 * cm

        for line in beschreibung.splitlines():
            c.drawString(x, y, line.strip())
            y -= 0.6 * cm
            if y < 4 * cm:
                c.showPage()
                y = A4[1] - 2 * cm

        # === Bild einfügen ===
        if "bild" in data:
            try:
                image_data = dh.read_binary(data["bild"])
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_img:
                    tmp_img.write(image_data)
                    tmp_img.flush()
                    c.drawImage(tmp_img.name, x, y - 6 * cm, width=10 * cm, height=6 * cm, preserveAspectRatio=True)
                    y -= 7 * cm
                    if y < 4 * cm:
                        c.showPage()
                        y = A4[1] - 2 * cm
            except Exception as e:
                c.drawString(x, y, f"⚠️ Bild {data.get('bild', 'unbekannt')} fehlerhaft: {str(e)}")
                y -= 1 * cm

        y -= 1 * cm  # Abstand zum nächsten Eintrag

    c.save()
    with open(tmp_pdf.name, "rb") as f:
        pdf_data = f.read()


    # === Download-Buttons anzeigen ===
    st.markdown("""
    <h2 style='text-decoration: underline; font-weight: bold; font-size: 28px; margin-top: 20px;'>
        Gesamten Zellatlas herunterladen
    </h2>
    """, unsafe_allow_html=True)

    if 'word_buffer' in locals():
        st.download_button("⬇️ Gesamtes Word-Dokument", data=word_buffer, file_name="Zellatlas_Haematologie.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    if 'pdf_data' in locals():
        st.download_button("⬇️ Gesamtes PDF-Dokument", data=pdf_data, file_name="Zellatlas_Haematologie.pdf", mime="application/pdf")

# === Zurück zur Übersicht ===
st.markdown("---")
if st.button("🔙 Zurück zur Übersicht"):
    st.switch_page("pages/01_Datei.py")
