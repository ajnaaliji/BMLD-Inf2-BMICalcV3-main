import streamlit as st
import pandas as pd
import io
import base64
from datetime import datetime
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
import tempfile
from PIL import Image
from docx.shared import Inches
import uuid
from utils.data_manager import DataManager
from utils.ui_helpers import apply_theme

# ==== Icon laden ====
def load_icon_base64(path):
    with open(path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode("utf-8")
    
img_chemie = load_icon_base64("assets/chemie.png")
img_paper = load_icon_base64("assets/paperclip.png")
img_pic = load_icon_base64("assets/picture.png")
img_datum = load_icon_base64("assets/calendar.png")
img_semester = load_icon_base64("assets/semester.png")
img_contract = load_icon_base64("assets/contract.png")
img_experiment = load_icon_base64("assets/experiment.png")
img_steps = load_icon_base64("assets/steps.png")
img_frage = load_icon_base64("assets/question.png")
img_ziel = load_icon_base64("assets/goal-flag.png")
img_chemical = load_icon_base64("assets/chemical.png")
img_tube = load_icon_base64("assets/test-tube.png")

def load_icon_bytes(path):
    with open(path, "rb") as image_file:
        return image_file.read()
    
icon_bytes = load_icon_bytes("assets/test-tube.png")

# Konvertiere zu PIL.Image (für Streamlit-kompatibles Format)
icon_image = Image.open(io.BytesIO(icon_bytes))

st.set_page_config(
    page_title="Chemie",
    page_icon=icon_image,
    layout="centered",
    initial_sidebar_state="collapsed"
)
# ==== Initialisierung ====
data_manager = DataManager()
username = st.session_state.get("username", "anonymous")
data_manager.load_user_data("chemie_eintraege", "data_chemie.csv", initial_value=[])

# ==== DataHandler vorbereiten ====
dh_word = data_manager._get_data_handler(f"word_chemie/{username}")
dh_img = data_manager._get_data_handler(f"bilder_chemie/{username}")
dh_docs = data_manager._get_data_handler(f"anhang_chemie/{username}")
for dh in [dh_word, dh_img, dh_docs]:
    if not dh.filesystem.exists(dh.root_path):
        dh.filesystem.makedirs(dh.root_path)

# ==== Titel ====
st.markdown(f"""
<h1 style='display: flex; align-items: center; gap: 24px;'>
    Chemie
    <img src='data:image/png;base64,{img_chemie}' width='50'>
</h1>
""", unsafe_allow_html=True)

# ==== Formular ====
col1, col2 = st.columns([2, 1])
with col1:
    with col1:
        st.markdown(f"""
        <div style="display: flex; align-items: center; gap: 10px;">
            <img src="data:image/png;base64,{img_experiment}" width="28">
            <h4 style="margin: 0;">Titel des Praktikums</h4>
        </div>
        """, unsafe_allow_html=True)
        titel = st.text_input("")
with col2:
    st.markdown(f"""
    <div style="display: flex; align-items: center; gap: 10px;">
        <img src="data:image/png;base64,{img_datum}" width="28">
        <h4 style="margin: 0;">Datum</h4>
    </div>
    """, unsafe_allow_html=True)
    datum = st.date_input("")
    st.markdown(f"""
    <div style="display: flex; align-items: center; gap: 10px;">
        <img src="data:image/png;base64,{img_semester}" width="28">
        <h4 style="margin: 0;">Semester</h4>
    </div>
    """, unsafe_allow_html=True)
    semester = st.selectbox("", ["1", "2", "3", "4", "5", "6"], key="semester")

st.markdown(f"""
<div style="display: flex; align-items: center; gap: 10px;">
    <img src="data:image/png;base64,{img_contract}" width="28">
    <h4 style="margin: 0;">Beschreibung des Versuchs</h4>
</div>
""", unsafe_allow_html=True)
beschreibung = st.text_area("", height=120, key="beschreibung")
col3, col4 = st.columns(2)
with col3:
    with col3:
        st.markdown(f"""
        <div style="display: flex; align-items: center; gap: 10px;">
            <img src="data:image/png;base64,{img_chemical}" width="28">
            <h4 style="margin: 0;">Benötigtes Material</h4>
        </div>
        """, unsafe_allow_html=True)
        material = st.text_area("", height=100, key="material")
with col4:
    with col4:
        # ==== Vorbereitung + Fragen ====
        st.markdown(f"""
        <h4 style='display: flex; align-items: center; gap: 10px; margin-top: 20px;'>
            <img src='data:image/png;base64,{img_frage}' width='28'>
            Vorbereitung + Fragen
        </h4>
        """, unsafe_allow_html=True)
        fragen = st.text_area("", height=100, key="fragen")
st.markdown(f"""
<div style="display: flex; align-items: center; gap: 10px;">
    <img src="data:image/png;base64,{img_steps}" width="28">
    <h4 style="margin: 0;">Arbeitsschritte</h4>
</div>
""", unsafe_allow_html=True)
arbeitsschritte = st.text_area("", height=120, key="arbeitsschritte")

st.markdown(f"""
<div style="display: flex; align-items: center; gap: 10px;">
    <img src="data:image/png;base64,{img_ziel}" width="28">
    <h4 style="margin: 0;">Ziel des Versuchs</h4>
</div>
""", unsafe_allow_html=True)
ziel = st.text_area("", height=100, key="ziel")

# ==== Bilder hochladen ====
st.markdown(f"""
<h3 style='display: flex; align-items: center; gap: 10px;'>
    <img src='data:image/png;base64,{img_pic}' width='40'>
    Mikroskopiebilder oder Befundbilder
</h3>
""", unsafe_allow_html=True)
uploaded_images = st.file_uploader("Wähle ein oder mehrere Bilder", type=["png", "jpg", "jpeg"], accept_multiple_files=True)
temp_uploaded_images = []

if uploaded_images:
    st.markdown("**Vorschau:**")
    bestehende_bilder = [f["name"] for f in dh_img.filesystem.ls(dh_img.root_path)]
    for img in uploaded_images:
        st.image(img, use_container_width=True)
        image_bytes = img.getvalue()
        clean_name = img.name.replace(" ", "_").replace("\u00e4", "ae").replace("\u00fc", "ue").replace("\u00f6", "oe")
        if any(f.endswith(clean_name) for f in bestehende_bilder):
            st.info(f"⏭️ Bild bereits vorhanden: {clean_name}")
            continue
        temp_uploaded_images.append((clean_name, image_bytes))

# ==== Anhänge ====
st.markdown(f"""
<h3 style='display: flex; align-items: center; gap: 10px;'>
    <img src='data:image/png;base64,{img_paper}' width='40'>
    Weitere Dateien (PDF, Word)
</h3>
""", unsafe_allow_html=True)
uploaded_docs = st.file_uploader("Wähle PDF- oder Word-Dokumente", type=["pdf", "docx"], accept_multiple_files=True)
temp_uploads = [(f.name, f.getvalue()) for f in uploaded_docs] if uploaded_docs else []
anhang_dateien = []

# ==== Speichern & Exportieren ====
if st.button("📎 Speichern und Exportieren"):
    if not titel.strip():
        st.warning("⚠️ Bitte gib einen Titel ein.")
        st.stop()

    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")

    # Bilder speichern
    bestehende_bilder = [f["name"] for f in dh_img.filesystem.ls(dh_img.root_path)]
    for name, img_bytes in temp_uploaded_images:
        if not any(f.endswith(name) for f in bestehende_bilder):
            dh_img.save(f"{timestamp}_{uuid.uuid4().hex}_{name}", img_bytes)

    # Anhänge speichern
    bestehende = [f["name"] for f in dh_docs.filesystem.ls(dh_docs.root_path)]
    for name, content in temp_uploads:
        name_clean = name.replace(" ", "_")
        if any(name_clean in f for f in bestehende):
            st.info(f"⏭️ Datei bereits vorhanden: {name}")
            continue
        filename = f"{timestamp}_{uuid.uuid4().hex[:8]}_{name_clean}"
        dh_docs.save(filename, content)
        anhang_dateien.append(filename)

    # ==== Word erstellen ====
    doc = Document()
    doc.add_heading(f"Praktikum: {titel}", 0)
    doc.add_paragraph(f"Datum: {datum.strftime('%d.%m.%Y')}")
    for header, content in [("Beschreibung", beschreibung), ("Material", material),
                            ("Vorbereitung + Fragen", fragen), ("Arbeitsschritte", arbeitsschritte), ("Ziel", ziel)]:
        doc.add_heading(header, level=2)
        doc.add_paragraph(content)

    if temp_uploaded_images:
        doc.add_page_break()
        doc.add_heading("Mikroskopiebilder / Versuchsbilder", level=2)
        for name, img_bytes in temp_uploaded_images:
            try:
                image_pil = Image.open(io.BytesIO(img_bytes))
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_img:
                    image_pil.save(tmp_img.name)
                    doc.add_picture(tmp_img.name, width=Inches(4.5))
            except Exception as e:
                st.warning(f"⚠️ Bild konnte nicht eingefügt werden: {name} ({e})")

    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    safe_title = titel.strip().replace(" ", "_").replace("/", "-").replace("\\", "-")
    filename_word = f"{timestamp}_{safe_title}.docx"
    dh_word.save(filename_word, word_buffer.getvalue())

    # ==== PDF erstellen ====
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
        c = canvas.Canvas(tmp_pdf.name, pagesize=A4)
        x, y = 2 * cm, A4[1] - 2 * cm
        c.setFont("Helvetica-Bold", 16)
        c.drawString(x, y, f"Praktikum: {titel}")
        y -= 1.5 * cm
        c.setFont("Helvetica", 12)
        c.drawString(x, y, f"Datum: {datum.strftime('%d.%m.%Y')}")
        y -= 1.5 * cm

        for header, content in [("Beschreibung", beschreibung), ("Material", material),
                                ("Vorbereitung + Fragen", fragen), ("Arbeitsschritte", arbeitsschritte), ("Ziel", ziel)]:
            c.setFont("Helvetica-Bold", 14)
            c.drawString(x, y, header)
            y -= 1 * cm
            c.setFont("Helvetica", 12)
            for line in content.splitlines():
                c.drawString(x, y, line.strip())
                y -= 0.6 * cm
                if y < 2 * cm:
                    c.showPage()
                    y = A4[1] - 2 * cm
        c.save()
        pdf_filename = f"{timestamp}_{uuid.uuid4().hex[:8]}_{safe_title}.pdf"
        with open(tmp_pdf.name, "rb") as f:
            pdf_bytes = f.read()
            dh_docs.save(pdf_filename, pdf_bytes)
            anhang_dateien.append(pdf_filename)

    # ==== Speichern als DataFrame ====
    neuer_eintrag = {
        "titel": titel,
        "datum": datum.strftime("%Y-%m-%d"),
        "beschreibung": beschreibung,
        "material": material,
        "fragen": fragen,
        "arbeitsschritte": arbeitsschritte,
        "ziel": ziel,
        "anhaenge": anhang_dateien,
        "semester": str(semester),
        "dateiname": filename_word,
        "zeit": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }

    df_neu = pd.DataFrame([neuer_eintrag])
    if isinstance(st.session_state["chemie_eintraege"], pd.DataFrame):
        st.session_state["chemie_eintraege"] = pd.concat([st.session_state["chemie_eintraege"], df_neu], ignore_index=True)
    else:
        st.session_state["chemie_eintraege"] = df_neu

    data_manager.save_data("chemie_eintraege")
    st.success("✅ Eintrag gespeichert!")

    st.download_button("⬇️ Word herunterladen", data=word_buffer, file_name=filename_word,
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    st.download_button("⬇️ PDF herunterladen", data=pdf_bytes, file_name=pdf_filename, mime="application/pdf")

# ==== Zurück zur Übersicht ====
if st.button("🔙 Zurück zur Übersicht"):
    st.switch_page("pages/01_Datei.py")
