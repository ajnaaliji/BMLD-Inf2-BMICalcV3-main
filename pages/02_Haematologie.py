import streamlit as st
import pandas as pd
import io
import base64
from datetime import datetime
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
import tempfile
from PIL import Image
from docx.shared import Inches
import uuid
from utils.data_manager import DataManager
from utils.ui_helpers import apply_theme
from PIL import Image, UnidentifiedImageError
import os

# ==== Initialisierung ====
st.set_page_config(page_title="Hämatologie", page_icon="🩸")
apply_theme()
data_manager = DataManager()
username = st.session_state.get("username", "anonymous")
data_manager.load_user_data("haematologie_eintraege", "data_haematologie.csv", initial_value=[])

# ==== Word & Anhang-Verzeichnisse vorbereiten ====
dh_word = data_manager._get_data_handler(f"word_haematologie/{username}")
dh_img = data_manager._get_data_handler(f"bilder_haematologie/{username}")
dh_docs = data_manager._get_data_handler(f"anhang_haematologie/{username}")
for dh in [dh_word, dh_img, dh_docs]:
    if not dh.filesystem.exists(dh.root_path):
        dh.filesystem.makedirs(dh.root_path)

# ==== Icon laden ====
def load_icon_base64(path):
    with open(path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode("utf-8")

img_rbb = load_icon_base64("assets/red-blood-cells.png")
img_neutro = load_icon_base64("assets/neutrophil.png")
img_lympho = load_icon_base64("assets/lymphocyte.png")
img_platelet = load_icon_base64("assets/platelet.png")
img_title = load_icon_base64("assets/blood-count.png")
img_blood = load_icon_base64("assets/blood.png")
img_paper = load_icon_base64("assets/paperclip.png")
img_pic = load_icon_base64("assets/picture.png")

# ==== Zellzählung ====
st.markdown(f"""
<h1 style='display: flex; align-items: center; gap: 24px;'>
    Hämatologie
    <img src='data:image/png;base64,{img_blood}' width='50'>
</h1>
""", unsafe_allow_html=True)

# ==== Eingabe: Titel, Datum, Befund ====
titel = st.text_input("Titel des Befundes")
datum = st.date_input("Datum", value=datetime.today())
semester = st.selectbox("Semester", ["1", "2", "3", "4", "5", "6"], key="semester")


zelltypen = [
    "Blasten", "Promyelozyten", "Myelozyten", "Metamyelozyten",
    "Stabkernige Granulozyten", "Segmentkernige Granulozyten",
    "Eosinophile", "Basophile", "Monozyten", "Lymphozyten", "Plasmazellen",
    "Erythroblasten", "Unbekannt / Diverses"
]

rb_felder = [
    "Anisozytose", "Mikrozyten", "Makrozyten", "Anisochromasie",
    "Hypochrom", "Hyperchrom", "Polychromasie", "Poikilozytose",
    "Ovalozyten", "Akanthozyten", "Sphärozyten", "Stomatozyten",
    "Echinozyten", "Targetzellen", "Tränenformen", "Sichelzellen",
    "Fragmentozyten", "Baso. Punktierung", "Howell Jollies", "Pappenheim",
]

gb_felder = [
    "vergröberte Granula", "basophile Schlieren",
    "Zytoplasmavakuolen", "Fehlende Granula", "Kernpyknose", "Pseudopelger",
    "Linksverschiebung", "Kerne hoch-/übersegmentiert"
]

ly_felder = [">10% LGL", "reaktiv", "pathologisch", "lymphoplasmozytoid"]
th_felder = ["Grosse Formen", "Riesenformen", "Agranulär"]

st.markdown("""
<style>
.zell-header {
    font-weight: bold;
    background-color: #f2f2f2;
    padding: 6px;
    border: 1px solid #ccc;
    text-align: center;
}
.zell-row {
    border-bottom: 1px solid #eee;
    padding: 8px 0;
}
.zell-cell {
    padding: 4px;
    text-align: center;
}
</style>
""", unsafe_allow_html=True)

# Header manuell
st.markdown("""
<div class="zell-row">
    <div class="zell-cell" style="display: inline-block; width: 20%;"><div class="zell-header">Zelltyp</div></div>
    <div class="zell-cell" style="display: inline-block; width: 30%;"><div class="zell-header">Zählung 1</div></div>
    <div class="zell-cell" style="display: inline-block; width: 30%;"><div class="zell-header">Zählung 2</div></div>
    <div class="zell-cell" style="display: inline-block; width: 15%;"><div class="zell-header">Durchschnitt</div></div>
</div>
""", unsafe_allow_html=True)

for zelltyp in zelltypen:
    z1_key = f"{zelltyp}_z1"
    z2_key = f"{zelltyp}_z2"

    if z1_key not in st.session_state:
        st.session_state[z1_key] = 0
    if z2_key not in st.session_state:
        st.session_state[z2_key] = 0

    col1, col2, col3, col4 = st.columns([2, 3, 3, 2])

    with col1:
        st.markdown(f"**{zelltyp}**")

    with col2:
        c1, c2, c3 = st.columns([1, 1, 1])
        if c1.button("➖", key=f"sub_{zelltyp}_z1"):
            st.session_state[z1_key] = max(0, st.session_state[z1_key] - 1)
        if c3.button("➕", key=f"add_{zelltyp}_z1"):
            st.session_state[z1_key] += 1
        c2.markdown(f"<div style='text-align: center; padding-top: 8px;'>{st.session_state[z1_key]}</div>", unsafe_allow_html=True)

    with col3:
        c4, c5, c6 = st.columns([1, 1, 1])
        if c4.button("➖", key=f"sub_{zelltyp}_z2"):
            st.session_state[z2_key] = max(0, st.session_state[z2_key] - 1)
        if c6.button("➕", key=f"add_{zelltyp}_z2"):
            st.session_state[z2_key] += 1
        c5.markdown(f"<div style='text-align: center; padding-top: 8px;'>{st.session_state[z2_key]}</div>", unsafe_allow_html=True)

    with col4:
        avg = (st.session_state[z1_key] + st.session_state[z2_key]) / 2
        st.markdown(f"**{avg:.1f}**")

st.markdown("---")

# ==== Zusatzbereiche: RBB, NG, LY, TH ====
def render_section(title, img_b64, felder, key_prefix, sonstiges_key):
    st.markdown(f"""
    <h3 style='display: flex; align-items: center; gap: 10px;'>
        {title}
        <img src='data:image/png;base64,{img_b64}' width='30'>
    </h3>
    """, unsafe_allow_html=True)
    for feld in felder:
        st.selectbox(feld, ["-", "+", "++", "+++"] , key=f"{key_prefix}_{feld}")
    st.text_area("Sonstiges:", key=sonstiges_key, height=80)

render_section("Rotes Blutbild", img_rbb, rb_felder, "rb", "rb_sonstiges")
render_section("Neutrophile Granulozyten", img_neutro, gb_felder, "gb", "ng_sonstiges")
render_section("Lymphozytenveränderungen", img_lympho, ly_felder, "ly", "lc_sonstiges")
render_section("Thrombozyten", img_platelet, th_felder, "th", "th_sonstiges")

# ==== Kontrolle auf fehlende Felder ====
fehlende = []
for feld in rb_felder:
    if st.session_state.get(f"rb_{feld}", "") == "":
        fehlende.append(f"Rotes Blutbild: {feld}")
for feld in gb_felder:
    if st.session_state.get(f"gb_{feld}", "") == "":
        fehlende.append(f"Granulozyten: {feld}")
for feld in ly_felder:
    if st.session_state.get(f"ly_{feld}", "") == "":
        fehlende.append(f"Lymphozyten: {feld}")
for feld in th_felder:
    if st.session_state.get(f"th_{feld}", "") == "":
        fehlende.append(f"Thrombozyten: {feld}")

if fehlende:
    st.error("❌ Bitte alle Felder ausfüllen:")
    for f in fehlende:
        st.markdown(f"- {f}")
    st.stop()

# ==== Bild-Upload ====
st.markdown(f"""
<h3 style='display: flex; align-items: center; gap: 10px;'>
    <img src='data:image/png;base64,{img_pic}' width='40'>
    Mikroskopiebilder oder Befundbilder
</h3>
""", unsafe_allow_html=True)
uploaded_images = st.file_uploader("Bilder auswählen", type=["png", "jpg", "jpeg"], accept_multiple_files=True)
temp_uploaded_images = []
if uploaded_images:
        st.markdown("**Vorschau:**")
        valide_uploaded_images = []  
        bestehende = [f["name"] for f in dh_img.filesystem.ls(dh_img.root_path)]
        for img in uploaded_images:
            st.image(img, use_container_width=True)
            name_clean = img.name.replace(" ", "_").replace("\u00e4", "ae").replace("\u00f6", "oe").replace("\u00fc", "ue")
            if any(f.endswith(name_clean) for f in bestehende):
                st.info(f"⏭️ Bild bereits vorhanden: {name_clean}")
                continue

            try:
                img_bytes = img.getvalue()
                image_pil = Image.open(io.BytesIO(img_bytes))
                st.text(f"{name_clean} – Größe: {image_pil.size}")
                if image_pil.size[0] == 0 or image_pil.size[1] == 0:
                    raise ValueError("Bildgröße ist 0")
                valide_uploaded_images.append((name_clean, img_bytes))
            except (UnidentifiedImageError, ValueError, ZeroDivisionError) as e:
                st.warning(f"⚠️ Bild konnte nicht eingefügt werden: {name_clean} ({e})")

# ==== Datei-Upload ====
st.markdown(f"""
<h3 style='display: flex; align-items: center; gap: 10px;'>
    <img src='data:image/png;base64,{img_paper}' width='40'>
    Weitere Dateien (PDF, Word)
</h3>
""", unsafe_allow_html=True)
uploaded_docs = st.file_uploader("Dateien auswählen", type=["pdf", "docx"], accept_multiple_files=True)
temp_uploads = [(f.name, f.getvalue()) for f in uploaded_docs] if uploaded_docs else []
anhang_dateien = []

# ==== Speichern & Exportieren ====
if "haema_exported" not in st.session_state:
    st.session_state["haema_exported"] = False

if st.button("📂 Speichern und Exportieren") and not st.session_state["haema_exported"]:
    st.session_state["haema_exported"] = True
    if not titel.strip():
        st.warning("Bitte einen Titel eingeben.")
        st.stop()

    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")

    # ==== Bilder speichern ====
    for name, img_bytes in valide_uploaded_images:
        dh_img.save(f"{timestamp}_{uuid.uuid4().hex}_{name}", img_bytes)

    # ==== Anhänge speichern ====
    for name, content in temp_uploads:
        name_clean = name.replace(" ", "_")
        filename = f"{timestamp}_{uuid.uuid4().hex[:8]}_{name_clean}"
        dh_docs.save(filename, content)
        anhang_dateien.append(filename)

    # ==== Word erstellen ====
    doc = Document()
    doc.add_heading(f"Befund: {titel}", 0)
    doc.add_paragraph(f"Datum: {datum.strftime('%d.%m.%Y')}")

    # Zellzählung erfassen
    doc.add_heading("Zellzählung", level=2)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Zelltyp"
    hdr_cells[1].text = "Zählung 1"
    hdr_cells[2].text = "Zählung 2"
    hdr_cells[3].text = "Durchschnitt"

    for zelltyp in zelltypen:
        z1 = st.session_state.get(f"{zelltyp}_z1", 0)
        z2 = st.session_state.get(f"{zelltyp}_z2", 0)
        avg = (z1 + z2) / 2
        row_cells = table.add_row().cells
        row_cells[0].text = zelltyp
        row_cells[1].text = str(z1)
        row_cells[2].text = str(z2)
        row_cells[3].text = f"{avg:.1f}"

    # Zusatzbereiche
    def add_section_to_doc(title, felder, prefix, sonstiges_key):
        doc.add_heading(title, level=2)
        for feld in felder:
            wert = st.session_state.get(f"{prefix}_{feld}", "-")
            doc.add_paragraph(f"{feld}: {wert}")
        doc.add_paragraph("Sonstiges:")
        doc.add_paragraph(st.session_state.get(sonstiges_key, "-"))

    add_section_to_doc("Rotes Blutbild", rb_felder, "rb", "rb_sonstiges")
    add_section_to_doc("Neutrophile Granulozyten", gb_felder, "gb", "ng_sonstiges")
    add_section_to_doc("Lymphozytenveränderungen", ly_felder, "ly", "lc_sonstiges")
    add_section_to_doc("Thrombozyten", th_felder, "th", "th_sonstiges")

    if valide_uploaded_images:
        doc.add_page_break()
        doc.add_heading("Bilder", level=2)

        for name, img_bytes in valide_uploaded_images:
            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_img_file:
                    tmp_img_file.write(img_bytes)
                    tmp_img_file_path = tmp_img_file.name
                doc.add_picture(tmp_img_file_path, width=Inches(4.5))
                doc.add_paragraph(name)
            except Exception as e:
                st.warning(f"⚠️ Bild konnte nicht eingefügt werden: {name} ({e})")
    from PIL import UnidentifiedImageError

    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    import re

    safe_titel = re.sub(r"[^\w\-_.]", "_", titel.strip())
    if not safe_titel:
        st.warning("Der Titel enthält keine gültigen Zeichen für einen Dateinamen.")
        st.stop()

    filename_word = f"{timestamp}_{safe_titel}.docx"
    dh_word.save(filename_word, word_buffer.getvalue())

    # ==== PDF erstellen ====
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_pdf:
        c = canvas.Canvas(tmp_pdf.name, pagesize=A4)
        x, y = 2 * cm, A4[1] - 2 * cm

    # Titel & Datum
        c.setFont("Helvetica-Bold", 16)
        c.drawString(x, y, f"Befund: {titel}")
        y -= 1.2 * cm
        c.setFont("Helvetica", 12)
        c.drawString(x, y, f"Datum: {datum.strftime('%d.%m.%Y')}")
        y -= 1.5 * cm

    # Zellzählung als Tabelle
        c.setFont("Helvetica-Bold", 14)
        c.drawString(x, y, "Zellzählung")
        y -= 1 * cm
        c.setFont("Helvetica-Bold", 12)
        c.drawString(x, y, "Zelltyp")
        c.drawString(x + 6 * cm, y, "Zählung 1")
        c.drawString(x + 9 * cm, y, "Zählung 2")
        c.drawString(x + 12 * cm, y, "Ø")
        y -= 0.7 * cm
        c.setFont("Helvetica", 12)

        for zelltyp in zelltypen:
            z1 = st.session_state.get(f"{zelltyp}_z1", 0)
            z2 = st.session_state.get(f"{zelltyp}_z2", 0)
            avg = (z1 + z2) / 2
            c.drawString(x, y, zelltyp)
            c.drawString(x + 6 * cm, y, str(z1))
            c.drawString(x + 9 * cm, y, str(z2))
            c.drawString(x + 12 * cm, y, f"{avg:.1f}")
            y -= 0.5 * cm
            if y < 3 * cm:
                c.showPage()
                y = A4[1] - 2 * cm

        # Zusatzbereiche wie im Word-Dokument
        def draw_section(c, x, y, title, felder, prefix, sonst_key):
            c.setFont("Helvetica-Bold", 14)
            c.drawString(x, y, title)
            y -= 0.8 * cm
            c.setFont("Helvetica", 12)
            for feld in felder:
                wert = st.session_state.get(f"{prefix}_{feld}", "-")
                c.drawString(x, y, f"{feld}: {wert}")
                y -= 0.5 * cm
                if y < 3 * cm:
                    c.showPage()
                    y = A4[1] - 2 * cm
            c.drawString(x, y, "Sonstiges:")
            y -= 0.5 * cm
            sonst_text = st.session_state.get(sonst_key, "-")
            for line in sonst_text.splitlines():
                c.drawString(x, y, line)
                y -= 0.5 * cm
                if y < 3 * cm:
                    c.showPage()
                    y = A4[1] - 2 * cm
            y -= 1 * cm
            return y

        y = draw_section(c, x, y, "Rotes Blutbild", rb_felder, "rb", "rb_sonstiges")
        y = draw_section(c, x, y, "Neutrophile Granulozyten", gb_felder, "gb", "ng_sonstiges")
        y = draw_section(c, x, y, "Lymphozytenveränderungen", ly_felder, "ly", "lc_sonstiges")
        y = draw_section(c, x, y, "Thrombozyten", th_felder, "th", "th_sonstiges")

    # Bilder einfügen (robustere Version mit PIL)

    if valide_uploaded_images:
        c.showPage()
        y = A4[1] - 2 * cm
        c.setFont("Helvetica-Bold", 14)
        c.drawString(x, y, "Bilder")
        y -= 1.2 * cm

        for name, img_bytes in valide_uploaded_images:
            try:
                # Bild öffnen mit PIL
                image = Image.open(io.BytesIO(img_bytes))
                image = image.convert("RGB")  # Für JPEG/PDF sicherstellen

                # Temporär als PNG speichern
                tmp_path = f"{tempfile.gettempdir()}/{uuid.uuid4().hex}.png"
                image.save(tmp_path, format="PNG")

                # Sicherheit: Breite/Höhe prüfen
                if image.width == 0 or image.height == 0:
                    raise ValueError("Bildgröße ist 0")

                # Bild ins PDF zeichnen
                c.drawImage(tmp_path, x, y - 7 * cm, width=12 * cm, height=6 * cm)
                y -= 8 * cm
                if y < 3 * cm:
                    c.showPage()
                    y = A4[1] - 2 * cm

            except Exception as e:
                st.warning(f"⚠️ Bild konnte nicht ins PDF eingefügt werden: {name} ({e})")
  
    c.save()

    # PDF-Datei speichern
    pdf_filename = f"{timestamp}_{safe_titel}.pdf"
    with open(tmp_pdf.name, "rb") as f:
        pdf_bytes = f.read()
    dh_docs.save(pdf_filename, pdf_bytes)
    anhang_dateien.append(pdf_filename)

    from PIL import UnidentifiedImageError

    # ==== Eintrag speichern ====
    neuer_eintrag = {
        "titel": titel,
        "datum": datum.strftime("%Y-%m-%d"),
        "anhaenge": anhang_dateien,
        "dateiname": filename_word,
        "pdfname": pdf_filename,   # <--- DAS FEHLT!
        "zeit": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }

    neuer_eintrag["semester"] = str(semester)

    df_neu = pd.DataFrame([neuer_eintrag])
    if isinstance(st.session_state["haematologie_eintraege"], pd.DataFrame):
        st.session_state["haematologie_eintraege"] = pd.concat([
            st.session_state["haematologie_eintraege"], df_neu
        ], ignore_index=True)
    else:
        st.session_state["haematologie_eintraege"] = df_neu

    data_manager.save_data("haematologie_eintraege")
    st.success("✅ Eintrag gespeichert!")

    st.download_button("⬇️ Word herunterladen", data=word_buffer, file_name=filename_word,
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    st.download_button("⬇️ PDF herunterladen", data=pdf_bytes, file_name=pdf_filename, mime="application/pdf")

# ==== Zurück ==== 
st.session_state["haema_exported"] = False

if st.button("🔙 Zurück zur Übersicht"):
    st.session_state["haema_exported"] = False
    st.switch_page("pages/01_Datei.py")
